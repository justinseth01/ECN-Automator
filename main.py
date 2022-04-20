# This is the main python script which sends the email, writes
# to the excel and word documents.
# It first calls our mainWindow.py file and displays the window.
# After the user closes the window, it will run the rest of the program
# if the user pressed "okay".
# It starts by loading the variables from the mainWindow class into local variables
# Next, it uses a try catch to run methods: sendEmail() toExcel() and newDoc()
# If the code fails at any point, it will undo changes made to the excel document.

"""
main.py

This project is designed to automate the documentation process 
by doing the following:
    -   Obtaining user's ECN information through a UI
    -   Reading, writing, and analyzing an excel sheet based on given
        user information
    -   Completing a preformatted word document with the ECN information
        entered by the user
    -   Sending an automated email with newly created document attached
All of these tasks originally needed to be executed manually, and was a 
rather tedious task due to its simplicity. This program simplifies this
process as much as possible by automating every step involved. Note this
project also contains a settings window so the user does not need to repeat
information every time the program is used.

This main python script responsible for calling all other classes and functions.
Is responsible for displaying the main window, as well as reorganizing the
data recieved from the window. It is responsible for most non UI processes such as
writing to the excel document, word document, and sending an automated email. 

Contains:
    -   sendEmail() takes the newly created word document and attached it to an
        email intended to be sent to other engineers to approve the change
    -   toExcel() takes the data recieved from the user via the main window
        and writes it to the "ECN Log" excel document
    -   newDoc() references the preformatted ECN word document "ECNFormat.docx" 
        and completes it with the information recieved from the user
    -   The program calls the main window and loads the data collected
        into local variables.


"""
import sys
import openpyxl
import base64
import sendgrid
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName, FileType, Disposition)
#import os

from mainWindow import mainWindow
from errorWindow import error

from PyQt5.QtWidgets import *
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date
from docx import Document

# Function that automatically sends an email


def sendEmail():
    message = Mail(
        from_email=settings[2],
        to_emails=settings[3],
        subject='ECN-'+str(ecnNumber),
        html_content='Dear Mr. Butler,<br><br>Attached is documentation for ECN-'+str(ecnNumber)+'. The document has also been placed in the ECN-ECR folder. Related drawings are in the "Released Drawings for Filing" folder.<br><br>Thanks,<br><br>'+settings[0].capitalize()+' '+settings[1].capitalize())

    # read word document via read binary base base 64 into ecnoded file variable
    with open(settings[6]+'/ECN-'+str(ecnNumber)+'.docx', 'rb') as f:
        data = f.read()
        f.close()
    encoded_file = base64.b64encode(data).decode()

    # add the attachment using the encoded file
    attachedFile = Attachment(
        FileContent(encoded_file),
        FileName('/ECN-'+str(ecnNumber)+'.docx'),
        FileType('application/docx'),
        Disposition('attachment')
    )
    message.attachment = attachedFile

    # API Key
    sg = sendgrid.SendGridAPIClient(
        'INSERT API KEY HERE')
    response = sg.send(message)

    # Print the results to ensure email was sent
    print(response.status_code, response.body, response.headers)

# Method writing information to the excel document


def toExcel(ecnNumber):
    # load workbook
    workbook = openpyxl.load_workbook(settings[4])
    sheet = workbook.active

    # find corresponding row with our ECN
    currentRow = 1
    cell = sheet.cell(row=currentRow, column=1)
    while(cell.value != ecnNumber):
        currentRow += 1
        cell = sheet.cell(row=currentRow, column=1)

    # change to string for sheet location
    currentRow = str(currentRow)

    # Write the information to the corresponding column
    sheet["B"+currentRow] = settings[0][0].upper()+". " + \
        settings[1].capitalize()
    sheet["C"+currentRow] = todayDate
    sheet["D"+currentRow] = projectNumber
    sheet["E"+currentRow] = serialNumber.upper()
    sheet["F"+currentRow] = machine.capitalize()
    sheet["G"+currentRow] = settings[0][0].upper()+". " + \
        settings[1].capitalize()
    sheet["H"+currentRow] = todayDate

    # save the file
    workbook.save(settings[4])


# Method for creating a new ECN Word Document
def newDoc():

    # Label document variable as the ECNFormat document
    document = Document(settings[5])

    # Add info in table 0
    table0 = document.tables[0]
    table0.cell(1, 0).text = str(ecnNumber)
    table0.cell(
        1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add corresponding info to table 1
    # Algorithm adding X to corresponding row/column
    table1 = document.tables[1]
    for i, change in enumerate(changes):
        if change.isChecked() == True:
            if i < window.rows:
                table1.cell(i, 1).text = "X"
            elif i < window.rows * 2:
                table1.cell(i-window.rows, 3).text = "X"
            else:
                table1.cell(i-2*window.rows, 5).text = "X"

    # Add corresponding info to table 2
    document.tables[2].cell(1, 0).text = description[0].upper()

    # Add corresponding info to table 3
    document.tables[3].cell(1, 0).text = str(mmNumber[0]) + \
        ' ('+partNumber[0].upper()+') \ REV '+revNo.upper()

    # Add corresponding info to table 4
    table4 = document.tables[4]
    table4.alignment
    for x in range(1, len(mmNumber)):
        table4.cell(x, 0).text = str(mmNumber[x])
        table4.cell(x, 1).text = partNumber[x].upper()
        table4.cell(x, 2).text = description[x].upper()
        table4.cell(x, 3).text = str(oldQuantity[x])
        table4.cell(x, 4).text = str(newQuantity[x])
        table4.cell(x, 5).text = str(statusCode[x])

    # Add corresponding info to table 5
    table5 = document.tables[5]
    table5.cell(1, 0).text = str(projectNumber)
    table5.cell(1, 1).text = serialNumber.upper()
    table5.cell(1, 2).text = str(machine)

    # Add corresponding info to table 6
    table6 = document.tables[6]
    table6.cell(1, 0).text = settings[0][0:1].upper(
    )+". "+settings[1].capitalize()
    table6.cell(1, 1).text = todayDate
    table6.cell(1, 2).text = notes

    document.save(settings[6]+'/ECN-'+str(ecnNumber)+'.docx')


app = QApplication(sys.argv)

# create the instance of our user interface window
window = mainWindow()

# showing the window
window.show()

# start the app
app.exec()

# Once the app is done running, then the rest of the program will run

# Only continue if the okay button was pressed and all error checks pass
# This is here in case the user presses the x button which closes the window. If this was not here,
# the program would run anyway
if not window.okay:
    print("ECN will not be made")
    # unreserve the ECN if not done so already
    window.unreserveECN(window.settings[4], int(window.ECN))
    sys.exit(0)

# make date/time variables for simplicity
today = date.today()
todayDate = today.strftime("%m/%d/%Y")

# use settings variable as an array of all our settings for simplicity
settings = window.settingsWindow.readSettings()
"""
settings[0]= first name
settings[1]= last name
settings[2]= from email
settings[3]= to email
settings[4]= ECN Log Path
settings[5]= ECN Format Path
settings[6]= ECN Save Path
"""

# Initialize our arrays that we will use
mmNumber = []
partNumber = []
description = []
oldQuantity = []
newQuantity = []
statusCode = []

# load our variables from the window class instance into local variables
ecnNumber = window.ECN
mmNumber.append(int(window.mmNumber.text()))
partNumber.append(window.partNumber.text())
revNo = window.revNumber.text()
description.append(window.description.text())
notes = window.notes.text()
projectNumber = int(window.projectNumber.text())
serialNumber = window.serialNumber.text()
machine = window.machine.text()
oldQuantity.append(0)  # No old QTY, new QTY or status code
newQuantity.append(0)  # for basic part info. Included to keep
statusCode.append(0)  # aligned with other array indicies
changes = window.changes

for button in window.Button1:  # load changes into variables
    mmNumber.append(button.mmNumber.text())
    partNumber.append(button.partNumber.text())
    description.append(button.description.text())
    oldQuantity.append(button.oldQuantity.text())
    newQuantity.append(button.newQuantity.text())
    statusCode.append(button.statusCode.text())

# Use try catch for writing to the excel document
try:
    toExcel(int(ecnNumber))
except Exception as e:
    # Show an error window then exit if program could not write to excel document
    error("Could not write to ECN Log. You must manually un-reserve the ECN", e)
    sys.exit(0)

# try catch for creating a new word document
try:
    newDoc()
except Exception as e:
    # show error window if word document could not be created
    error("Could not create new ECN Word Document", e)

    # undo changes in excel document then quit program
    window.unreserveECN(settings[4], int(ecnNumber))
    sys.exit(0)

# try catch to send email
try:
    sendEmail()
except Exception as e:
    # show error window if word document could not be created
    error("Could not send email", e)

    # undo changes in excel document then quit program
    window.unreserveECN(settings[4], int(ecnNumber))
    sys.exit(0)
